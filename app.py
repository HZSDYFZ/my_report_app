# -*- coding: utf-8 -*-
import io
import re
import zipfile
import pandas as pd
import streamlit as st
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from datetime import datetime

st.set_page_config(page_title="认证评定报告自动化生成系统", page_icon="📄", layout="wide")

BOX_PATTERN = re.compile(r"^[□☐☑✔\[\]口\s]+")

def extract_first_person(lead_str):
    """提取审核组长姓名"""
    if pd.isna(lead_str) or not str(lead_str).strip():
        return ""
    s = str(lead_str).strip()
    s = re.sub(r"^(审核组长|组长|Lead)[:：]\s*", "", s, flags=re.IGNORECASE)
    s = re.sub(r"[\（\(].*?[\）\)]", "", s).strip()
    parts = re.split(r"[ ,，/／、+&\t\n]+", s)
    return parts[0] if parts and parts[0] else ""

def clean_date_val(val):
    """将 Excel 日期统一转换为标准 YYYY-MM-DD 格式"""
    if pd.isna(val):
        return ""
    if isinstance(val, (pd.Timestamp, datetime)):
        return val.strftime("%Y-%m-%d")
    val_str = str(val).strip()
    if not val_str or val_str.lower() in ["nan", "none", "null", "nat", "0", "undefined"]:
        return ""
    
    match = re.search(r'(\d{2,4})[-–‑—](\d{1,2})[-–‑—](\d{1,2})', val_str)
    if match:
        groups = match.groups()
        year = groups[0]
        if len(year) == 2:
            year = "20" + year
        month = groups[1].zfill(2)
        day = groups[2].zfill(2)
        return f"{year}-{month}-{day}"

    try:
        f_val = float(val_str)
        if 30000 < f_val < 60000:
            dt = pd.to_datetime(f_val, unit='D', origin='1899-12-30')
            return dt.strftime("%Y-%m-%d")
    except ValueError:
        pass
    
    try:
        dt = pd.to_datetime(val_str)
        if not pd.isna(dt):
            return dt.strftime("%Y-%m-%d")
    except Exception:
        pass
        
    return val_str.split(" ")[0]

def format_chinese_date(date_str):
    """将 YYYY-MM-DD 转换为 某年某月某日 格式"""
    if not date_str:
        return ""
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return f"{dt.year}年{dt.month}月{dt.day}日"
    except Exception:
        return date_str

def find_best_column(columns, possible_keys):
    """智能匹配表头"""
    for key in possible_keys:
        for col in columns:
            col_clean = str(col).replace(" ", "").replace("\n", "").lower()
            key_clean = key.replace(" ", "").lower()
            if key_clean in col_clean:
                return col
    return columns[0] if len(columns) > 0 else None

def remove_invisible_chars(text):
    """清除 Word 文本中常见的零宽空格等隐藏控制字符"""
    if not text:
        return ""
    return re.sub(r'[\u200b\u200c\u200d\u00ad\ufeff]', '', text)

def process_paragraph_text(p, data):
    """合并段落内所有 Run 进行全局替换，彻底解决 Word 拆分标签导致的匹配失败问题"""
    try:
        full_text = "".join([run.text for run in p.runs])
        if not full_text.strip():
            return
    except Exception:
        return

    chinese_date = format_chinese_date(data["eval_date"])

    replacements = {
        "{{公司名称}}": data["company_name"],
        "{{ 公司名称 }}": data["company_name"],
        "{{任务号}}": data["task_no"],
        "{{ 任务号 }}": data["task_no"],
        "{{审核组长}}": data["lead"],
        "{{ 审核组长 }}": data["lead"],
        "{{审核地址}}": data["address"],
        "{{ 审核地址 }}": data["address"],
        "{{审核范围}}": data["scope"],
        "{{ 审核范围 }}": data["scope"],
        "{{评定日期}}": chinese_date,
        "{{ 评定日期 }}": chinese_date,
        "{{评定通过时间}}": chinese_date,
        "{{ 评定通过时间 }}": chinese_date,
        "【公司名称】": data["company_name"],
        "【任务号】": data["task_no"],
        "【审核组长】": data["lead"],
        "【审核地址】": data["address"],
        "【审核范围】": data["scope"],
        "【评定日期】": chinese_date,
        "【评定通过时间】": chinese_date,
        
        # 认证标准复选框
        "{{iatf_check}}": "☑" if data["has_ts"] else "☐",
        "{{ iatf_check }}": "☑" if data["has_ts"] else "☐",
        "{{iso_check}}": "☑" if data["has_er"] else "☐",
        "{{ iso_check }}": "☑" if data["has_er"] else "☐",

        # 审核类型复选框
        "{{chu_shen}}": "☑" if data["is_initial"] else "☐",
        "{{ chu_shen }}": "☑" if data["is_initial"] else "☐",
        "{{jian_shen}}": "☑" if data["is_surveillance"] else "☐",
        "{{ jian_shen }}": "☑" if data["is_surveillance"] else "☐",
        "{{zai_ren_zheng}}": "☑" if data["is_recert_transfer"] else "☐",
        "{{ zai_ren_zheng }}": "☑" if data["is_recert_transfer"] else "☐",
        "{{te_shu}}": "☑" if data["is_special"] else "☐",
        "{{ te_shu }}": "☑" if data["is_special"] else "☐",

        # 认证决定结论复选框
        "{{dec_1}}": "☑" if data["decision_option"] == 1 else "☐",
        "{{ dec_1 }}": "☑" if data["decision_option"] == 1 else "☐",
        "{{dec_3}}": "☑" if data["decision_option"] == 3 else "☐",
        "{{ dec_3 }}": "☑" if data["decision_option"] == 3 else "☐",
        "{{dec_4}}": "☑" if data["decision_option"] == 4 else "☐",
        "{{ dec_4 }}": "☑" if data["decision_option"] == 4 else "☐",
        "{{dec_5}}": "☑" if data["decision_option"] == 5 else "☐",
        "{{ dec_5 }}": "☑" if data["decision_option"] == 5 else "☐",
        "{{dec_6}}": "☑" if data["decision_option"] == 6 else "☐",
        "{{ dec_6 }}": "☑" if data["decision_option"] == 6 else "☐",
    }

    modified = False
    new_text = remove_invisible_chars(full_text)

    # 1. 替换所有标签字典
    for k, v in replacements.items():
        if k in new_text:
            new_text = new_text.replace(k, str(v))
            modified = True

    # 2. 精准将日期写在“日期：”后面
    if chinese_date:
        for prefix in ["评定日期", "评审日期", "决定日期", "日期", "评定通过时间"]:
            if prefix in new_text:
                pattern = r'(' + prefix + r'\s*[：:])\s*([_—\s]*|\d{4}[年\-\./]\d{1,2}[月\-\./]\d{1,2}日?)'
                if re.search(pattern, new_text):
                    new_text = re.sub(pattern, r'\1 ' + chinese_date, new_text)
                    modified = True

    if modified and p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""

def fill_next_target_cell(cells, current_idx, value):
    """基于底层 XML 单元格定位标签格后面紧挨着的下一个独立单元格"""
    if not str(value).strip():
        return
    current_tc = cells[current_idx]._tc
    for next_idx in range(current_idx + 1, len(cells)):
        if cells[next_idx]._tc != current_tc:
            target_cell = cells[next_idx]
            if target_cell.paragraphs:
                target_cell.paragraphs[0].text = str(value)
            else:
                target_cell.add_paragraph(str(value))
            break

def format_decision_options(cell, data):
    """还原认证决定结论格式，支持多选项精准匹配勾选"""
    option_paragraphs = [p for p in cell.paragraphs if "通过" in p.text or "不予通过" in p.text or BOX_PATTERN.search(p.text)]
    target_idx = data["decision_option"]

    for opt_i, p in enumerate(option_paragraphs, start=1):
        raw_text = p.text.strip()
        clean_text = BOX_PATTERN.sub("", raw_text).strip()

        is_selected = (opt_i == target_idx)
        mark = "☑ " if is_selected else "☐ "

        if opt_i == 1 and "适用于：" in clean_text:
            prefix = clean_text.split("适用于：")[0]
            clean_text = f"{prefix}适用于：{data['scope']}）"

        p.text = mark + clean_text

def process_table_safely(table, data):
    """表格安全遍历"""
    visited_tcs = set()

    for row in table.rows:
        cells = row.cells
        for idx, cell in enumerate(cells):
            if cell._tc in visited_tcs:
                continue
            visited_tcs.add(cell._tc)

            cell_text = remove_invisible_chars(cell.text.strip())
            clean_text = re.sub(r"[\s:：]", "", cell_text)

            for p in cell.paragraphs:
                process_paragraph_text(p, data)

            if clean_text in ["公司名称", "客户名称", "企业名称"]:
                if cell.paragraphs:
                    cell.paragraphs[0].text = f"公司名称：{data['company_name']}"

            elif clean_text in ["任务号", "合同号"]:
                if cell.paragraphs:
                    cell.paragraphs[0].text = f"任务号：{data['task_no']}"

            elif (clean_text in ["审核组长", "组长"]) and "报告评定人员" not in clean_text:
                fill_next_target_cell(cells, idx, data['lead'])

            for p in cell.paragraphs:
                p_clean = re.sub(r"\s+", "", remove_invisible_chars(p.text))
                if "审核地址：" in p_clean or "审核地址:" in p_clean:
                    p.text = f"审核地址：{data['address']}"
                elif "认证范围：" in p_clean or "认证范围:" in p_clean:
                    p.text = f"认证范围：{data['scope']}"

            if clean_text in ["日期", "评定日期", "评定通过时间", "评定时间", "通过日期"]:
                chinese_date = format_chinese_date(data["eval_date"])
                if cell.paragraphs and chinese_date:
                    for p in cell.paragraphs:
                        process_paragraph_text(p, data)

            if "认证决定结论" in clean_text or ("通过" in cell_text and "不予通过" in cell_text):
                format_decision_options(cell, data)

def fill_word_template(template_bytes, data):
    doc = Document(io.BytesIO(template_bytes))

    def process_container(container):
        for p in container.paragraphs:
            process_paragraph_text(p, data)
        for table in container.tables:
            process_table_safely(table, data)
        for p_elem in container._element.xpath('.//w:txbxContent//w:p'):
            p = Paragraph(p_elem, doc)
            process_paragraph_text(p, data)
        for tbl_elem in container._element.xpath('.//w:txbxContent//w:tbl'):
            table = Table(tbl_elem, doc)
            process_table_safely(table, data)

    process_container(doc)

    for section in doc.sections:
        hf_list = [
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer
        ]
        for hf in hf_list:
            try:
                process_container(hf)
            except Exception:
                pass

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream.getvalue()

# Streamlit 界面
st.title("📄 认证评定报告自动化生成系统")

c1, c2 = st.columns(2)
with c1:
    excel_file = st.file_uploader("1. 上传认证 Excel 数据文件 (.xlsx / .xls)", type=["xlsx", "xls"])
with c2:
    template_file = st.file_uploader("2. 上传 Word 报告模板 (.docx)", type=["docx"])

st.markdown("---")

if excel_file is not None and template_file is not None:
    try:
        raw_df = pd.read_excel(excel_file)
        template_bytes = template_file.getvalue()
        columns = list(raw_df.columns)

        # 后台自动智能匹配列名
        col_comp = find_best_column(columns, ["公司名称", "客户名称", "企业名称", "公司"])
        col_task = find_best_column(columns, ["任务号", "合同号", "项目编号"])
        col_lead = find_best_column(columns, ["审核组长", "组长", "审核员"])
        col_date = find_best_column(columns, ["评定通过时间", "评定日期", "决定日期", "日期", "评审日期", "时间"])
        col_addr = find_best_column(columns, ["审核地址", "地址"])
        col_scope = find_best_column(columns, ["审核范围", "认证范围", "范围"])
        col_type = find_best_column(columns, ["审核类型", "类型", "阶段"])
        col_conclusion = find_best_column(columns, ["认证决定结论", "决定结论", "结论"])

        parsed_records = []
        for idx, row in raw_df.iterrows():
            comp_val = str(row.get(col_comp, "")) if col_comp else ""
            task_val = str(row.get(col_task, "")) if col_task else ""
            lead_val = str(row.get(col_lead, "")) if col_lead else ""
            
            if pd.isna(row.dropna()).all() or (comp_val.lower() in ["nan", "none", "", "0"] and task_val.lower() in ["nan", "none", "", "0"]):
                continue

            company_name = comp_val if comp_val.lower() not in ["nan", "none"] else ""
            task_no = task_val if task_val.lower() not in ["nan", "none"] else ""
            lead_first = extract_first_person(lead_val)
            address = str(row.get(col_addr, "")) if col_addr and not pd.isna(row.get(col_addr)) else ""
            scope = str(row.get(col_scope, "")) if col_scope and not pd.isna(row.get(col_scope)) else ""
            audit_type_raw = str(row.get(col_type, "")) if col_type and not pd.isna(row.get(col_type)) else ""
            decision_conclusion = str(row.get(col_conclusion, "")) if col_conclusion and not pd.isna(row.get(col_conclusion)) else ""
            
            eval_date_raw = row.get(col_date, "") if col_date else ""
            eval_date = clean_date_val(eval_date_raw)

            task_upper = task_no.upper()
            
            # 标准勾选判定：兼容任务号或审核类型中的标识
            has_ts = "TS" in task_upper or "16949" in task_upper or "TS" in audit_type_raw.upper()
            has_er = "ER" in task_upper or "9001" in task_upper or "9001" in audit_type_raw

            # 审核类型勾选判定（增强兼容性）
            is_initial = "二阶段" in audit_type_raw or "初审" in audit_type_raw or "第1阶段" in audit_type_raw
            is_surveillance = "监" in audit_type_raw or "监督" in audit_type_raw
            is_recert_transfer = "再认证" in audit_type_raw or "转换" in audit_type_raw or "转移" in audit_type_raw
            is_special = "特殊" in audit_type_raw

            # 认证决定结论选项判定逻辑
            decision_option = 1  # 默认第一行
            if "二阶段" in audit_type_raw or "初审" in audit_type_raw or "再认证" in audit_type_raw:
                decision_option = 1
            elif "转移" in audit_type_raw or "转换" in audit_type_raw:
                decision_option = 3
            elif ("监" in audit_type_raw or "监督" in audit_type_raw) and "不换证" in decision_conclusion:
                decision_option = 5
            elif ("监" in audit_type_raw or "监督" in audit_type_raw) and "换发" in decision_conclusion:
                decision_option = 6
            elif "特殊" in audit_type_raw and "换发" in decision_conclusion:
                decision_option = 4

            parsed_records.append({
                "company_name": company_name,
                "task_no": task_no,
                "lead": lead_first,
                "address": address,
                "scope": scope,
                "has_ts": has_ts,
                "has_er": has_er,
                "is_initial": is_initial,
                "is_surveillance": is_surveillance,
                "is_recert_transfer": is_recert_transfer,
                "is_special": is_special,
                "audit_type_raw": audit_type_raw,
                "decision_conclusion": decision_conclusion,
                "decision_option": decision_option,
                "eval_date": eval_date
            })

        if not parsed_records:
            st.warning("⚠️ Excel 文件中未读取到有效数据。")
        else:
            st.subheader(f"📋 预览数据（共 {len(parsed_records)} 条）")
            preview_df = pd.DataFrame(parsed_records)[
                ["company_name", "task_no", "lead", "audit_type_raw", "decision_conclusion", "decision_option", "eval_date"]
            ]
            preview_df.columns = ["公司名称", "任务号", "审核组长", "审核类型", "决定结论", "勾选结论选项", "评定日期"]
            st.dataframe(preview_df, use_container_width=True)

            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for idx, data in enumerate(parsed_records):
                    doc_bytes = fill_word_template(template_bytes, data)
                    clean_company = re.sub(r'[\\/*?:"<>|]', "_", str(data["company_name"]) or "未命名企业")
                    clean_task = re.sub(r'[\\/*?:"<>|]', "_", str(data["task_no"]) or "未命名任务")
                    zf.writestr(f"{clean_company}_{clean_task}_评定报告.docx", doc_bytes)

            zip_buffer.seek(0)

            st.download_button(
                label=f"📦 下载批量生成报告包 ({len(parsed_records)} 份 .zip)",
                data=zip_buffer.getvalue(),
                file_name="认证评定报告包.zip",
                mime="application/zip",
                type="primary",
                use_container_width=True
            )

    except Exception as e:
        st.error(f"处理失败: {str(e)}")
else:
    st.info("👈 请上传对应的 Excel 和 Word 模板文件进行处理。")
